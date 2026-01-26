"""
Media Plan Document Generator Azure Function
Version: 5.3
Purpose: Generate Word documents programmatically with dynamic row counts for DMAs, channels, audiences

This function receives JSON payload and builds a complete media plan document using python-docx.
Unlike template-based approaches, this handles variable row counts dynamically - only populated
data appears in the final document.

Dynamic Sections:
- Budget Allocation: 1-5 tiers
- Channel Mix: 3-8 channels  
- DMA Allocation: 4-10 DMAs per tier
- Performance Table: 5-20 total DMAs
- Audience Insights: 2-6 segments
- Behavioral Signals: 4-12 types
- Contextual Categories: 4-10 categories
"""

import azure.functions as func
import json
import logging
from io import BytesIO
from datetime import datetime
from typing import Dict, List, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Color Constants (matching Output_Templates_v5_1.txt)
class Colors:
    PRIMARY_BLUE = RGBColor(30, 58, 95)      # #1E3A5F
    SECONDARY_BLUE = RGBColor(44, 82, 130)   # #2C5282
    ACCENT_BLUE = RGBColor(49, 130, 206)     # #3182CE
    LIGHT_BLUE = RGBColor(235, 248, 255)     # #EBF8FF
    BLACK = RGBColor(26, 32, 44)             # #1A202C
    DARK_GRAY = RGBColor(45, 55, 72)         # #2D3748
    MEDIUM_GRAY = RGBColor(113, 128, 150)    # #718096
    LIGHT_GRAY = RGBColor(226, 232, 240)     # #E2E8F0
    BACKGROUND_GRAY = RGBColor(247, 250, 252) # #F7FAFC
    WHITE = RGBColor(255, 255, 255)          # #FFFFFF
    SUCCESS_GREEN = RGBColor(56, 161, 105)   # #38A169
    WARNING_YELLOW = RGBColor(214, 158, 46)  # #D69E2E
    ALERT_ORANGE = RGBColor(221, 107, 32)    # #DD6B20
    ERROR_RED = RGBColor(229, 62, 62)        # #E53E3E


class DocumentBuilder:
    """Builds Word documents with professional formatting and dynamic content."""
    
    def __init__(self, data: Dict[str, Any]):
        self.data = data
        self.doc = Document()
        self._setup_document_styles()
        self._setup_page_layout()
    
    def _setup_document_styles(self):
        """Configure document-wide styles."""
        styles = self.doc.styles
        
        # Default paragraph style
        style_normal = styles['Normal']
        style_normal.font.name = 'Arial'
        style_normal.font.size = Pt(11)
        style_normal.font.color.rgb = Colors.BLACK
        style_normal.paragraph_format.space_after = Pt(6)
        style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style_normal.paragraph_format.line_spacing = 1.15
        
        # Heading 1 style
        style_h1 = styles['Heading 1']
        style_h1.font.name = 'Arial'
        style_h1.font.size = Pt(18)
        style_h1.font.bold = True
        style_h1.font.color.rgb = Colors.PRIMARY_BLUE
        style_h1.paragraph_format.space_before = Pt(12)
        style_h1.paragraph_format.space_after = Pt(12)
        
        # Heading 2 style
        style_h2 = styles['Heading 2']
        style_h2.font.name = 'Arial'
        style_h2.font.size = Pt(14)
        style_h2.font.bold = True
        style_h2.font.color.rgb = Colors.SECONDARY_BLUE
        style_h2.paragraph_format.space_before = Pt(10)
        style_h2.paragraph_format.space_after = Pt(10)
        
        # Heading 3 style
        style_h3 = styles['Heading 3']
        style_h3.font.name = 'Arial'
        style_h3.font.size = Pt(12)
        style_h3.font.bold = True
        style_h3.font.color.rgb = Colors.DARK_GRAY
        style_h3.paragraph_format.space_before = Pt(8)
        style_h3.paragraph_format.space_after = Pt(8)
        
        # Create custom Title style
        try:
            style_title = styles.add_style('DocumentTitle', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style_title = styles['DocumentTitle']
        style_title.font.name = 'Arial'
        style_title.font.size = Pt(28)
        style_title.font.bold = True
        style_title.font.color.rgb = Colors.PRIMARY_BLUE
        style_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_title.paragraph_format.space_after = Pt(18)
        
        # Create Caption style
        try:
            style_caption = styles.add_style('TableCaption', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style_caption = styles['TableCaption']
        style_caption.font.name = 'Arial'
        style_caption.font.size = Pt(9)
        style_caption.font.italic = True
        style_caption.font.color.rgb = Colors.MEDIUM_GRAY
        style_caption.paragraph_format.space_after = Pt(4)
    
    def _setup_page_layout(self):
        """Configure page margins and layout."""
        sections = self.doc.sections
        for section in sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
    
    def _set_cell_shading(self, cell, color: RGBColor):
        """Apply background shading to a table cell."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), f'{color.red:02X}{color.green:02X}{color.blue:02X}')
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _set_cell_borders(self, cell, color: RGBColor = Colors.LIGHT_GRAY, size: int = 4):
        """Apply borders to a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(size))
            border.set(qn('w:color'), f'{color.red:02X}{color.green:02X}{color.blue:02X}')
            tcBorders.append(border)
        tcPr.append(tcBorders)
    
    def _create_table(self, headers: List[str], rows: List[List[str]], 
                      col_widths: Optional[List[float]] = None) -> None:
        """Create a formatted table with headers and data rows."""
        if not rows:
            return
        
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths if provided
        if col_widths:
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = Inches(width)
        
        # Header row
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = Colors.WHITE
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            self._set_cell_shading(cell, Colors.PRIMARY_BLUE)
            self._set_cell_borders(cell)
        
        # Data rows
        for row_idx, row_data in enumerate(rows):
            data_row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                cell = data_row.cells[col_idx]
                cell.text = str(cell_text) if cell_text is not None else ''
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].runs[0].font.color.rgb = Colors.BLACK
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._set_cell_borders(cell)
                
                # Alternate row shading
                if row_idx % 2 == 1:
                    self._set_cell_shading(cell, Colors.BACKGROUND_GRAY)
        
        self.doc.add_paragraph()
    
    def _create_kpi_table(self, kpis: List[Dict]) -> None:
        """Create a KPI table with status indicators."""
        if not kpis:
            return
        
        headers = ['KPI', 'Target', 'Projection', 'Confidence', 'Status']
        rows = []
        for kpi in kpis:
            status = kpi.get('status', 'On Track')
            rows.append([
                kpi.get('name', ''),
                kpi.get('target', ''),
                kpi.get('projection', ''),
                kpi.get('confidence', ''),
                status
            ])
        
        self._create_table(headers, rows, [2.0, 1.3, 1.3, 1.5, 1.2])

    def _add_cover_page(self) -> None:
        """Generate the cover page."""
        metadata = self.data.get('metadata', {})
        
        # Add spacing at top
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Main title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('MEDIA PLAN')
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = Colors.PRIMARY_BLUE
        
        self.doc.add_paragraph()
        
        # Campaign name
        campaign = self.doc.add_paragraph()
        campaign.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = campaign.add_run(metadata.get('campaignName', 'Campaign Name'))
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = Colors.SECONDARY_BLUE
        
        # Campaign period
        period = self.doc.add_paragraph()
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = period.add_run(metadata.get('campaignPeriod', 'Campaign Period'))
        run.font.size = Pt(14)
        run.font.color.rgb = Colors.DARK_GRAY
        
        # Add spacing
        for _ in range(8):
            self.doc.add_paragraph()
        
        # Document info
        info_items = [
            ('Prepared for:', metadata.get('clientName', 'Client Name')),
            ('Prepared by:', metadata.get('preparedBy', 'Agency/Team Name')),
            ('Date:', metadata.get('documentDate', datetime.now().strftime('%B %d, %Y'))),
            ('Version:', metadata.get('version', '1.0'))
        ]
        
        for label, value in info_items:
            para = self.doc.add_paragraph()
            run_label = para.add_run(label + ' ')
            run_label.font.size = Pt(11)
            run_label.font.color.rgb = Colors.DARK_GRAY
            run_value = para.add_run(value)
            run_value.font.size = Pt(11)
            run_value.font.color.rgb = Colors.BLACK
        
        # Confidential footer
        for _ in range(4):
            self.doc.add_paragraph()
        
        conf = self.doc.add_paragraph()
        conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = conf.add_run('CONFIDENTIAL')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = Colors.MEDIUM_GRAY
        
        self.doc.add_page_break()
    
    def _add_executive_summary(self) -> None:
        """Generate the Executive Summary section."""
        summary = self.data.get('executiveSummary', {})
        metadata = self.data.get('metadata', {})
        
        self.doc.add_heading('EXECUTIVE SUMMARY', level=1)
        
        # Campaign At-A-Glance table
        self.doc.add_heading('Campaign At-A-Glance', level=2)
        
        glance_data = [
            ['Campaign Name', metadata.get('campaignName', '')],
            ['Client', metadata.get('clientName', '')],
            ['Campaign Period', metadata.get('campaignPeriod', '')],
            ['Total Budget', summary.get('totalBudget', '')],
            ['Primary Objective', summary.get('primaryObjective', '')],
            ['Primary KPI', summary.get('primaryKPI', '')],
            ['Geographic Focus', summary.get('geographicFocus', '')],
            ['Pathway Used', metadata.get('pathway', '')]
        ]
        
        self._create_table(['Element', 'Details'], glance_data, [2.5, 4.8])
        
        # Strategic Overview
        self.doc.add_heading('Strategic Overview', level=2)
        overview = summary.get('strategicOverview', '')
        if overview:
            para = self.doc.add_paragraph(overview)
            para.paragraph_format.space_after = Pt(12)
        
        # Budget Allocation Summary (DYNAMIC - 1-5 tiers)
        budget_allocation = self.data.get('budgetAllocation', {})
        tiers = budget_allocation.get('tiers', [])
        
        if tiers:
            self.doc.add_heading('Budget Allocation Summary', level=2)
            headers = ['Tier', 'Allocation %', 'Budget', 'Projected Performance']
            rows = []
            for tier in tiers:
                rows.append([
                    tier.get('name', ''),
                    tier.get('allocationPercent', ''),
                    tier.get('budget', ''),
                    tier.get('projectedPerformance', '')
                ])
            
            # Add totals row if present
            if budget_allocation.get('totalBudget'):
                rows.append(['TOTAL', '100%', budget_allocation.get('totalBudget', ''), ''])
            
            self._create_table(headers, rows, [2.0, 1.5, 1.8, 2.0])
        
        # KPI Targets and Projections
        kpis = summary.get('kpiTargets', [])
        if kpis:
            self.doc.add_heading('KPI Targets and Projections', level=2)
            self._create_kpi_table(kpis)
        
        # Key Recommendations
        recommendations = summary.get('keyRecommendations', [])
        if recommendations:
            self.doc.add_heading('Key Recommendations', level=2)
            for i, rec in enumerate(recommendations, 1):
                para = self.doc.add_paragraph()
                run_num = para.add_run(f'{i}. ')
                run_num.font.bold = True
                para.add_run(rec)
        
        self.doc.add_page_break()

    def _add_channel_strategy(self) -> None:
        """Generate the Channel Strategy section (DYNAMIC - 3-8 channels)."""
        channel_mix = self.data.get('channelMix', {})
        channels = channel_mix.get('channels', [])
        
        if not channels:
            return
        
        self.doc.add_heading('CHANNEL STRATEGY', level=1)
        
        # Channel Selection Rationale
        self.doc.add_heading('Channel Selection Rationale', level=2)
        
        headers = ['Channel', 'Primary Role', 'Funnel Stage', 'Selection Rationale']
        rows = []
        for channel in channels:
            rows.append([
                channel.get('name', ''),
                channel.get('primaryRole', ''),
                channel.get('funnelStage', ''),
                channel.get('selectionRationale', '')
            ])
        
        self._create_table(headers, rows, [1.8, 1.8, 1.5, 2.2])
        
        # Channel Allocation
        self.doc.add_heading('Channel Allocation', level=2)
        
        headers = ['Channel', 'Budget', '% of Total', 'Primary KPI', 'Projected CPM']
        rows = []
        for channel in channels:
            rows.append([
                channel.get('name', ''),
                channel.get('budget', ''),
                channel.get('percentOfTotal', ''),
                channel.get('primaryKPI', ''),
                channel.get('projectedCPM', '')
            ])
        
        self._create_table(headers, rows, [1.5, 1.4, 1.2, 1.5, 1.5])
        
        # Individual Channel Details
        self.doc.add_heading('Channel Details', level=2)
        
        for channel in channels:
            self.doc.add_heading(channel.get('name', 'Channel'), level=3)
            
            # Channel overview table
            overview_data = [
                ['Budget', channel.get('budget', '')],
                ['Flight Dates', channel.get('flightDates', '')],
                ['Primary Objective', channel.get('primaryObjective', '')],
                ['Ad Formats', channel.get('adFormats', '')],
                ['Targeting', channel.get('targeting', '')],
                ['Platforms', channel.get('platforms', '')]
            ]
            
            self._create_table(['Attribute', 'Details'], overview_data, [2.0, 5.3])
            
            # Performance Projections for this channel
            projections = channel.get('performanceProjections', [])
            if projections:
                para = self.doc.add_paragraph()
                run = para.add_run('Performance Projections')
                run.font.bold = True
                run.font.size = Pt(11)
                
                headers = ['Metric', 'Projection', 'Confidence', 'Benchmark']
                rows = []
                for proj in projections:
                    rows.append([
                        proj.get('metric', ''),
                        proj.get('projection', ''),
                        proj.get('confidence', ''),
                        proj.get('benchmark', '')
                    ])
                
                self._create_table(headers, rows, [1.8, 1.8, 1.5, 2.2])
        
        self.doc.add_page_break()
    
    def _add_dma_allocation(self) -> None:
        """Generate DMA Allocation section (DYNAMIC - 4-10 DMAs per tier, 5-20 total)."""
        dma_data = self.data.get('dmaAllocation', {})
        dma_tiers = dma_data.get('tiers', [])
        
        if not dma_tiers:
            return
        
        self.doc.add_heading('DMA ALLOCATION', level=1)
        
        # Overview
        overview = dma_data.get('overview', '')
        if overview:
            para = self.doc.add_paragraph(overview)
            para.paragraph_format.space_after = Pt(12)
        
        # DMA by Tier (each tier has 4-10 DMAs)
        for tier in dma_tiers:
            tier_name = tier.get('tierName', 'Tier')
            tier_budget = tier.get('tierBudget', '')
            tier_percent = tier.get('tierPercent', '')
            
            self.doc.add_heading(f'{tier_name} Markets ({tier_percent} - {tier_budget})', level=2)
            
            dmas = tier.get('dmas', [])
            if dmas:
                headers = ['DMA', 'Market', 'Population', 'Budget', '% of Tier', 'CPM Index']
                rows = []
                for dma in dmas:
                    rows.append([
                        dma.get('dmaCode', ''),
                        dma.get('marketName', ''),
                        dma.get('population', ''),
                        dma.get('budget', ''),
                        dma.get('percentOfTier', ''),
                        dma.get('cpmIndex', '')
                    ])
                
                self._create_table(headers, rows, [1.0, 2.0, 1.3, 1.3, 1.2, 1.2])
        
        # Performance Summary Table (all DMAs)
        all_dmas = dma_data.get('performanceSummary', [])
        if all_dmas:
            self.doc.add_heading('DMA Performance Summary', level=2)
            
            headers = ['DMA', 'Market', 'Tier', 'Budget', 'Projected Impressions', 'Projected CTR']
            rows = []
            for dma in all_dmas:
                rows.append([
                    dma.get('dmaCode', ''),
                    dma.get('marketName', ''),
                    dma.get('tier', ''),
                    dma.get('budget', ''),
                    dma.get('projectedImpressions', ''),
                    dma.get('projectedCTR', '')
                ])
            
            self._create_table(headers, rows, [1.0, 1.8, 1.0, 1.2, 1.5, 1.2])
        
        self.doc.add_page_break()

    def _add_target_audience(self) -> None:
        """Generate Target Audience section (DYNAMIC - 2-6 segments)."""
        audience_data = self.data.get('targetAudience', {})
        segments = audience_data.get('segments', [])
        
        if not segments:
            return
        
        self.doc.add_heading('TARGET AUDIENCE', level=1)
        
        for i, segment in enumerate(segments):
            priority = 'Primary' if i == 0 else f'Secondary {i}'
            segment_name = segment.get('name', 'Audience Segment')
            
            self.doc.add_heading(f'{priority} Audience: {segment_name}', level=2)
            
            # Demographic Profile
            demographics = segment.get('demographics', {})
            if demographics:
                para = self.doc.add_paragraph()
                run = para.add_run('Demographic Profile')
                run.font.bold = True
                
                demo_data = [
                    ['Age Range', demographics.get('ageRange', '')],
                    ['Gender', demographics.get('gender', '')],
                    ['Household Income', demographics.get('householdIncome', '')],
                    ['Education', demographics.get('education', '')],
                    ['Geography', demographics.get('geography', '')],
                    ['Urban/Suburban/Rural', demographics.get('urbanSuburbanRural', '')]
                ]
                
                self._create_table(['Attribute', 'Value'], demo_data, [2.5, 4.8])
            
            # Psychographic Profile
            psychographics = segment.get('psychographics', {})
            if psychographics:
                para = self.doc.add_paragraph()
                run = para.add_run('Psychographic Profile')
                run.font.bold = True
                
                psych_data = [
                    ['Interests', psychographics.get('interests', '')],
                    ['Values', psychographics.get('values', '')],
                    ['Lifestyle', psychographics.get('lifestyle', '')],
                    ['Media Consumption', psychographics.get('mediaConsumption', '')],
                    ['Purchase Behavior', psychographics.get('purchaseBehavior', '')]
                ]
                
                self._create_table(['Attribute', 'Characteristics'], psych_data, [2.0, 5.3])
            
            # Behavioral Signals (DYNAMIC - 4-12 types)
            behavioral_signals = segment.get('behavioralSignals', [])
            if behavioral_signals:
                para = self.doc.add_paragraph()
                run = para.add_run('Behavioral Indicators')
                run.font.bold = True
                
                headers = ['Behavior Type', 'Indicators', 'Data Source']
                rows = []
                for signal in behavioral_signals:
                    rows.append([
                        signal.get('behaviorType', ''),
                        signal.get('indicators', ''),
                        signal.get('dataSource', '')
                    ])
                
                self._create_table(headers, rows, [2.0, 3.0, 2.3])
        
        # Audience Insights
        insights = audience_data.get('insights', [])
        if insights:
            self.doc.add_heading('Audience Insights', level=2)
            
            headers = ['Source', 'Type', 'Key Finding']
            rows = []
            for insight in insights:
                rows.append([
                    insight.get('source', ''),
                    insight.get('type', ''),
                    insight.get('keyFinding', '')
                ])
            
            self._create_table(headers, rows, [2.0, 1.5, 3.8])
        
        self.doc.add_page_break()
    
    def _add_contextual_targeting(self) -> None:
        """Generate Contextual Targeting section (DYNAMIC - 4-10 categories)."""
        contextual = self.data.get('contextualTargeting', {})
        categories = contextual.get('categories', [])
        
        if not categories:
            return
        
        self.doc.add_heading('CONTEXTUAL TARGETING STRATEGY', level=1)
        
        # Overview
        overview = contextual.get('overview', '')
        if overview:
            para = self.doc.add_paragraph(overview)
            para.paragraph_format.space_after = Pt(12)
        
        # Categories table
        headers = ['Category', 'Priority', 'Channels', 'Budget Allocation', 'Expected Performance']
        rows = []
        for category in categories:
            rows.append([
                category.get('name', ''),
                category.get('priority', ''),
                category.get('channels', ''),
                category.get('budgetAllocation', ''),
                category.get('expectedPerformance', '')
            ])
        
        self._create_table(headers, rows, [2.0, 1.0, 1.5, 1.5, 1.5])
        
        # Exclusions
        exclusions = contextual.get('exclusions', [])
        if exclusions:
            self.doc.add_heading('Content Exclusions', level=2)
            
            headers = ['Exclusion Type', 'Specification', 'Rationale']
            rows = []
            for exclusion in exclusions:
                rows.append([
                    exclusion.get('type', ''),
                    exclusion.get('specification', ''),
                    exclusion.get('rationale', '')
                ])
            
            self._create_table(headers, rows, [2.0, 2.5, 2.8])
        
        self.doc.add_page_break()

    def _add_performance_projections(self) -> None:
        """Generate Performance Projections section."""
        projections = self.data.get('performanceProjections', {})
        
        if not projections:
            return
        
        self.doc.add_heading('PERFORMANCE PROJECTIONS', level=1)
        
        # Consolidated Projections
        funnel_projections = projections.get('funnelProjections', [])
        if funnel_projections:
            self.doc.add_heading('Funnel Projections', level=2)
            
            headers = ['Funnel Stage', 'Metric', 'Projection', 'Confidence', 'Range']
            rows = []
            for proj in funnel_projections:
                rows.append([
                    proj.get('funnelStage', ''),
                    proj.get('metric', ''),
                    proj.get('projection', ''),
                    proj.get('confidence', ''),
                    proj.get('range', '')
                ])
            
            self._create_table(headers, rows, [1.3, 1.5, 1.5, 1.5, 1.5])
        
        # Weekly Projections
        weekly_projections = projections.get('weeklyProjections', [])
        if weekly_projections:
            self.doc.add_heading('Weekly Projections', level=2)
            
            headers = ['Week', 'Spend', 'Impressions', 'Clicks', 'Conversions', 'Cumulative Conv.']
            rows = []
            for week in weekly_projections:
                rows.append([
                    week.get('week', ''),
                    week.get('spend', ''),
                    week.get('impressions', ''),
                    week.get('clicks', ''),
                    week.get('conversions', ''),
                    week.get('cumulativeConversions', '')
                ])
            
            self._create_table(headers, rows, [1.0, 1.3, 1.5, 1.2, 1.3, 1.5])
        
        # Scenario Analysis
        scenarios = projections.get('scenarios', {})
        if scenarios:
            self.doc.add_heading('Scenario Analysis', level=2)
            
            for scenario_name, scenario_data in scenarios.items():
                para = self.doc.add_paragraph()
                run = para.add_run(f'{scenario_name} Scenario')
                run.font.bold = True
                
                if isinstance(scenario_data, list):
                    headers = ['Metric', 'Value', 'vs Target']
                    rows = []
                    for metric in scenario_data:
                        rows.append([
                            metric.get('metric', ''),
                            metric.get('value', ''),
                            metric.get('vsTarget', '')
                        ])
                    
                    self._create_table(headers, rows, [2.5, 2.0, 2.0])
        
        self.doc.add_page_break()
    
    def _add_measurement_attribution(self) -> None:
        """Generate Measurement and Attribution section."""
        measurement = self.data.get('measurement', {})
        
        if not measurement:
            return
        
        self.doc.add_heading('MEASUREMENT AND ATTRIBUTION', level=1)
        
        # Primary Metrics
        primary_metrics = measurement.get('primaryMetrics', [])
        if primary_metrics:
            self.doc.add_heading('Primary Metrics', level=2)
            
            headers = ['Category', 'Metric', 'Target', 'Measurement Source', 'Frequency']
            rows = []
            for metric in primary_metrics:
                rows.append([
                    metric.get('category', ''),
                    metric.get('metric', ''),
                    metric.get('target', ''),
                    metric.get('measurementSource', ''),
                    metric.get('frequency', '')
                ])
            
            self._create_table(headers, rows, [1.3, 1.5, 1.3, 1.8, 1.3])
        
        # Attribution Model
        attribution = measurement.get('attributionModel', {})
        if attribution:
            self.doc.add_heading('Attribution Model', level=2)
            
            para = self.doc.add_paragraph()
            run = para.add_run(f"Selected Model: ")
            run.font.bold = True
            para.add_run(attribution.get('selectedModel', ''))
            
            parameters = attribution.get('parameters', [])
            if parameters:
                headers = ['Parameter', 'Setting', 'Rationale']
                rows = []
                for param in parameters:
                    rows.append([
                        param.get('parameter', ''),
                        param.get('setting', ''),
                        param.get('rationale', '')
                    ])
                
                self._create_table(headers, rows, [2.0, 2.0, 3.3])
        
        # Reporting Cadence
        reporting = measurement.get('reportingCadence', [])
        if reporting:
            self.doc.add_heading('Reporting Cadence', level=2)
            
            headers = ['Report Type', 'Frequency', 'Audience', 'Delivery Method']
            rows = []
            for report in reporting:
                rows.append([
                    report.get('reportType', ''),
                    report.get('frequency', ''),
                    report.get('audience', ''),
                    report.get('deliveryMethod', '')
                ])
            
            self._create_table(headers, rows, [2.0, 1.5, 2.0, 1.8])
        
        self.doc.add_page_break()

    def _add_appendices(self) -> None:
        """Generate Appendices section."""
        appendices = self.data.get('appendices', {})
        
        if not appendices:
            return
        
        self.doc.add_heading('APPENDICES', level=1)
        
        # Creative Asset Matrix
        creative_assets = appendices.get('creativeAssetMatrix', [])
        if creative_assets:
            self.doc.add_heading('Appendix A: Creative Asset Matrix', level=2)
            
            headers = ['Asset ID', 'Name', 'Format', 'Dimensions', 'Channel', 'Status']
            rows = []
            for asset in creative_assets:
                rows.append([
                    asset.get('assetId', ''),
                    asset.get('name', ''),
                    asset.get('format', ''),
                    asset.get('dimensions', ''),
                    asset.get('channel', ''),
                    asset.get('status', '')
                ])
            
            self._create_table(headers, rows, [1.0, 2.0, 1.0, 1.2, 1.3, 1.0])
        
        # DMA Details
        dma_details = appendices.get('dmaDetails', [])
        if dma_details:
            self.doc.add_heading('Appendix B: DMA Allocation Details', level=2)
            
            headers = ['DMA', 'Market', 'Population', 'Budget', '% Total']
            rows = []
            for dma in dma_details:
                rows.append([
                    dma.get('dmaCode', ''),
                    dma.get('marketName', ''),
                    dma.get('population', ''),
                    dma.get('budget', ''),
                    dma.get('percentTotal', '')
                ])
            
            self._create_table(headers, rows, [1.0, 2.5, 1.5, 1.5, 1.0])
        
        # Data Sources
        data_sources = appendices.get('dataSources', [])
        if data_sources:
            self.doc.add_heading('Appendix C: Data Sources', level=2)
            
            headers = ['Source', 'Type', 'Usage', 'Access']
            rows = []
            for source in data_sources:
                rows.append([
                    source.get('source', ''),
                    source.get('type', ''),
                    source.get('usage', ''),
                    source.get('access', '')
                ])
            
            self._create_table(headers, rows, [2.0, 1.5, 2.5, 1.5])
        
        # Glossary
        glossary = appendices.get('glossary', [])
        if glossary:
            self.doc.add_heading('Appendix D: Glossary of Terms', level=2)
            
            headers = ['Term', 'Definition']
            rows = []
            for term in glossary:
                rows.append([
                    term.get('term', ''),
                    term.get('definition', '')
                ])
            
            self._create_table(headers, rows, [2.0, 5.3])
        
        # Version History
        version_history = appendices.get('versionHistory', [])
        if version_history:
            self.doc.add_heading('Appendix E: Version History', level=2)
            
            headers = ['Version', 'Date', 'Author', 'Changes']
            rows = []
            for version in version_history:
                rows.append([
                    version.get('version', ''),
                    version.get('date', ''),
                    version.get('author', ''),
                    version.get('changes', '')
                ])
            
            self._create_table(headers, rows, [1.0, 1.5, 2.0, 3.0])
    
    def build(self) -> bytes:
        """Build the complete document and return as bytes."""
        try:
            self._add_cover_page()
            self._add_executive_summary()
            self._add_channel_strategy()
            self._add_dma_allocation()
            self._add_target_audience()
            self._add_contextual_targeting()
            self._add_performance_projections()
            self._add_measurement_attribution()
            self._add_appendices()
            
            # Save to bytes
            buffer = BytesIO()
            self.doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logging.error(f"Error building document: {str(e)}")
            raise


def validate_request(data: Dict[str, Any]) -> List[str]:
    """Validate the incoming request payload."""
    errors = []
    
    # Required top-level fields
    required_fields = ['metadata']
    for field in required_fields:
        if field not in data:
            errors.append(f"Missing required field: {field}")
    
    # Validate metadata
    metadata = data.get('metadata', {})
    required_metadata = ['campaignName', 'clientName']
    for field in required_metadata:
        if field not in metadata:
            errors.append(f"Missing required metadata field: {field}")
    
    return errors


def main(req: func.HttpRequest) -> func.HttpResponse:
    """
    HTTP Trigger handler for Media Plan document generation.
    
    Accepts POST requests with JSON payload containing media plan data.
    Returns a .docx file as binary response.
    """
    logging.info('Media Plan Document Generator function triggered.')
    
    try:
        # Parse request body
        try:
            request_body = req.get_json()
        except ValueError:
            return func.HttpResponse(
                json.dumps({
                    "status": "error",
                    "message": "Invalid JSON in request body",
                    "code": "INVALID_JSON"
                }),
                status_code=400,
                mimetype="application/json"
            )
        
        # Validate request
        validation_errors = validate_request(request_body)
        if validation_errors:
            return func.HttpResponse(
                json.dumps({
                    "status": "error",
                    "message": "Validation failed",
                    "errors": validation_errors,
                    "code": "VALIDATION_ERROR"
                }),
                status_code=400,
                mimetype="application/json"
            )
        
        # Build document
        builder = DocumentBuilder(request_body)
        doc_bytes = builder.build()
        
        # Generate filename
        metadata = request_body.get('metadata', {})
        campaign_name = metadata.get('campaignName', 'MediaPlan').replace(' ', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{campaign_name}_MediaPlan_{timestamp}.docx"
        
        logging.info(f'Successfully generated document: {filename}')
        
        # Return document as binary response
        return func.HttpResponse(
            body=doc_bytes,
            status_code=200,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(len(doc_bytes))
            }
        )
        
    except Exception as e:
        logging.error(f"Error generating document: {str(e)}")
        return func.HttpResponse(
            json.dumps({
                "status": "error",
                "message": str(e),
                "code": "GENERATION_ERROR"
            }),
            status_code=500,
            mimetype="application/json"
        )
