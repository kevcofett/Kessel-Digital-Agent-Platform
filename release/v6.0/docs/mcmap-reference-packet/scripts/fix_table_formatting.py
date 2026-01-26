#!/usr/bin/env python3
"""
Post-process DOCX files to add professional table formatting.
Pandoc generates tables with an unresolved 'Table' style that renders
as borderless in Word. This script adds explicit borders, header shading,
and consistent formatting to all tables.
"""

import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Twips, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """Set cell border properties.
    Usage: set_cell_border(cell, top={"sz": 8, "color": "404040", "val": "single"},
                                  bottom={"sz": 8, "color": "404040", "val": "single"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr_name, attr_val in kwargs[edge].items():
                element.set(qn(f'w:{attr_name}'), str(attr_val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders(table):
    """Set table-level borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # Remove any existing borders
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '8')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '404040')
        borders.append(element)
    tblPr.append(borders)


def set_cell_shading(cell, color_hex):
    """Set background shading for a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shading
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_margins(table):
    """Set consistent cell margins for the entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # Remove existing cell margins
    for existing in tblPr.findall(qn('w:tblCellMar')):
        tblPr.remove(existing)

    margins = OxmlElement('w:tblCellMar')
    for edge, width in [('top', '40'), ('left', '80'), ('bottom', '40'), ('right', '80')]:
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:w'), width)
        element.set(qn('w:type'), 'dxa')
        margins.append(element)
    tblPr.append(margins)


def format_table(table):
    """Apply full professional formatting to a table."""
    # 1. Set table-level borders
    set_table_borders(table)

    # 2. Set cell margins
    set_cell_margins(table)

    # 3. Format header row (first row)
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, 'F2F2F2')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # 4. Alternating row shading for data rows
    for i, row in enumerate(table.rows):
        if i == 0:
            continue  # Skip header
        color = 'FFFFFF' if i % 2 == 1 else 'FAFAFA'
        for cell in row.cells:
            set_cell_shading(cell, color)


def fix_tables_in_docx(docx_path):
    """Fix all tables in a DOCX file."""
    doc = Document(docx_path)
    table_count = len(doc.tables)

    for table in doc.tables:
        format_table(table)

    doc.save(docx_path)
    return table_count


def main():
    if len(sys.argv) > 1:
        docx_files = sys.argv[1:]
    else:
        base_dir = os.path.dirname(os.path.dirname(__file__))
        docx_dir = os.path.join(base_dir, 'docx')
        docx_files = [
            os.path.join(docx_dir, f)
            for f in sorted(os.listdir(docx_dir))
            if f.endswith('.docx') and not f.startswith('~')
        ]

    print("Fixing table formatting in DOCX files...")
    print(f"Processing {len(docx_files)} files\n")

    total_tables = 0
    for filepath in docx_files:
        filename = os.path.basename(filepath)
        try:
            count = fix_tables_in_docx(filepath)
            total_tables += count
            print(f"  {filename}: {count} tables formatted")
        except Exception as e:
            print(f"  {filename}: ERROR - {e}")

    print(f"\nComplete: {total_tables} tables formatted across {len(docx_files)} files")


if __name__ == '__main__':
    main()
