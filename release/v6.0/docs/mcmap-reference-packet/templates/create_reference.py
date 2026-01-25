#!/usr/bin/env python3
"""
Create reference.docx template for MCMAP documentation.
Footer contains:
  - Center: MASTERCARD CONFIDENTIAL (C) 2026
  - Right: Page number
"""

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

def add_page_number(paragraph):
    """Add a page number field to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    return run

def add_tab_stops(paragraph, center_pos, right_pos):
    """Add tab stops to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')

    # Center tab
    tab_center = OxmlElement('w:tab')
    tab_center.set(qn('w:val'), 'center')
    tab_center.set(qn('w:pos'), str(int(center_pos.twips)))
    tabs.append(tab_center)

    # Right tab
    tab_right = OxmlElement('w:tab')
    tab_right.set(qn('w:val'), 'right')
    tab_right.set(qn('w:pos'), str(int(right_pos.twips)))
    tabs.append(tab_right)

    pPr.append(tabs)

def create_reference_doc():
    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Enable footer
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear default paragraph and create new one
    footer_para = footer.paragraphs[0]
    footer_para.clear()

    # Add tab stops for center (3.25") and right (6.5") positions
    add_tab_stops(footer_para, Inches(3.25), Inches(6.5))

    # Add tab to center, then confidentiality text
    run_tab1 = footer_para.add_run('\t')
    run_center = footer_para.add_run('MASTERCARD CONFIDENTIAL ')
    run_center.font.size = Pt(9)
    run_center.font.name = 'Calibri'

    # Add copyright symbol and year
    run_copy = footer_para.add_run('\u00A9 2026')
    run_copy.font.size = Pt(9)
    run_copy.font.name = 'Calibri'

    # Add tab to right, then page number
    run_tab2 = footer_para.add_run('\t')
    run_page_label = footer_para.add_run('Page ')
    run_page_label.font.size = Pt(9)
    run_page_label.font.name = 'Calibri'

    page_run = add_page_number(footer_para)
    page_run.font.size = Pt(9)
    page_run.font.name = 'Calibri'

    # Add a sample heading to establish styles
    heading = doc.add_heading('Sample Heading 1', level=1)
    heading.runs[0].font.name = 'Calibri Light'
    heading.runs[0].font.size = Pt(16)

    heading2 = doc.add_heading('Sample Heading 2', level=2)
    heading2.runs[0].font.name = 'Calibri Light'
    heading2.runs[0].font.size = Pt(14)

    para = doc.add_paragraph('Sample body text for reference document.')
    para.runs[0].font.name = 'Calibri'
    para.runs[0].font.size = Pt(11)

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'reference.docx')
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

if __name__ == '__main__':
    create_reference_doc()
