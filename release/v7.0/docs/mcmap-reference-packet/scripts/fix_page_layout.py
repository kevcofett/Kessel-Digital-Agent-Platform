#!/usr/bin/env python3
"""
Post-process DOCX files for intelligent page layout.

Pandoc generates explicit page break paragraphs from {=openxml} raw blocks.
This script replaces that approach with style-based pagination:

1. Removes all explicit page break paragraphs (empty Normal paragraphs with w:br page breaks)
2. Adds pageBreakBefore to Heading 2 paragraphs (major sections always start on new pages)
3. Adds keepWithNext to all heading paragraphs (prevents orphaned headings)
4. Adds keepLines to all heading paragraphs (prevents heading text from splitting)

This produces cleaner pagination: major sections get new pages, subsections flow
naturally, and no more than ~1/3 of a page is left blank.
"""

import os
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADING_STYLES = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
                  'Heading 5', 'Heading 6', 'Title', 'Subtitle'}

PAGE_BREAK_BEFORE_STYLES = {'Heading 2'}


def is_explicit_page_break_paragraph(paragraph):
    """Check if a paragraph is an explicit page break (empty paragraph with only w:br type=page)."""
    p = paragraph._p

    # Must have runs
    runs = p.findall(qn('w:r'))
    if not runs:
        return False

    # Check if all runs contain only page break elements
    has_page_break = False
    has_text = False

    for run in runs:
        for child in run:
            tag = child.tag
            if tag == qn('w:br'):
                br_type = child.get(qn('w:type'), '')
                if br_type == 'page':
                    has_page_break = True
                else:
                    has_text = True
            elif tag == qn('w:t'):
                text = child.text or ''
                if text.strip():
                    has_text = True
            elif tag == qn('w:rPr'):
                pass  # Run properties are OK
            else:
                has_text = True

    return has_page_break and not has_text


def remove_explicit_page_breaks(doc):
    """Remove all explicit page break paragraphs from the document body."""
    removed = 0
    body = doc.element.body

    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        if is_explicit_page_break_paragraph(paragraph):
            paragraphs_to_remove.append(paragraph)

    for paragraph in paragraphs_to_remove:
        body.remove(paragraph._p)
        removed += 1

    return removed


def add_paragraph_property(paragraph, prop_name):
    """Add a paragraph formatting property (e.g., keepNext, keepLines, pageBreakBefore)."""
    pPr = paragraph._p.get_or_add_pPr()

    # Remove existing property if present
    for existing in pPr.findall(qn(f'w:{prop_name}')):
        pPr.remove(existing)

    element = OxmlElement(f'w:{prop_name}')
    pPr.append(element)


def fix_page_layout(docx_path):
    """Apply intelligent page layout to a DOCX file."""
    doc = Document(docx_path)

    # Step 1: Remove explicit page break paragraphs
    breaks_removed = remove_explicit_page_breaks(doc)

    # Step 2: Add pageBreakBefore to Heading 2 paragraphs
    # Skip if it's the very first paragraph in the document
    page_breaks_added = 0
    first_paragraph = True
    for paragraph in doc.paragraphs:
        if paragraph.style is None:
            continue

        style_name = paragraph.style.name
        if style_name is None:
            continue

        if first_paragraph:
            first_paragraph = False
            if style_name in PAGE_BREAK_BEFORE_STYLES:
                # Don't add page break before the very first heading
                continue

        if style_name in PAGE_BREAK_BEFORE_STYLES:
            add_paragraph_property(paragraph, 'pageBreakBefore')
            page_breaks_added += 1

    # Step 3: Add keepWithNext and keepLines to all heading paragraphs
    headings_formatted = 0
    for paragraph in doc.paragraphs:
        if paragraph.style is None:
            continue

        style_name = paragraph.style.name
        if style_name is None:
            continue

        if style_name in HEADING_STYLES:
            add_paragraph_property(paragraph, 'keepNext')
            add_paragraph_property(paragraph, 'keepLines')
            headings_formatted += 1

    doc.save(docx_path)
    return breaks_removed, page_breaks_added, headings_formatted


def main():
    if len(sys.argv) > 1:
        docx_files = sys.argv[1:]
    else:
        base_dir = os.path.dirname(os.path.dirname(__file__))
        docx_dir = os.path.join(base_dir, 'docx')
        docx_files = [
            os.path.join(docx_dir, f)
            for f in sorted(os.listdir(docx_dir))
            if f.endswith('.docx') and not f.startswith('~')
        ]

    print("Fixing page layout in DOCX files...")
    print(f"Processing {len(docx_files)} files\n")

    total_removed = 0
    total_breaks = 0
    total_headings = 0

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        try:
            removed, breaks, headings = fix_page_layout(filepath)
            total_removed += removed
            total_breaks += breaks
            total_headings += headings
            print(f"  {filename}: {removed} explicit breaks removed, "
                  f"{breaks} pageBreakBefore added, {headings} headings formatted")
        except Exception as e:
            print(f"  {filename}: ERROR - {e}")

    print(f"\nComplete: {total_removed} breaks removed, {total_breaks} pageBreakBefore added, "
          f"{total_headings} headings formatted across {len(docx_files)} files")


if __name__ == '__main__':
    main()
